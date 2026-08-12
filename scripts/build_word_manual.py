#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Gerador de Manual Oficial em Word (.docx) — Correio Digital Angola (CDA 2026)
Converte o MANUAL_DE_UTILIZACAO_CDA_2026.md em um documento .docx com formatação
profissional, tabelas estilizadas, hierarquia de cabeçalhos e caixas de destaque.
"""

import os
import re
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_table_borders(table, color="D3D3D3", sz="4", val="single"):
    tblPr = table._tbl.tblPr
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:bottom w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:insideH w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:insideV w:val="none"/>'
        f'<w:left w:val="none"/>'
        f'<w:right w:val="none"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)

def add_formatted_text(paragraph, text, default_font_name="Calibri", default_font_size=Pt(11), default_color=RGBColor(51, 65, 85)):
    # Simple markdown inline parser for **bold**, *italic*, `code`
    tokens = re.split(r'(\*\*.*?\*\*|\*.*?\*|`.*?`)', text)
    for token in tokens:
        if not token:
            continue
        run = paragraph.add_run()
        run.font.name = default_font_name
        run.font.size = default_font_size
        run.font.color.rgb = default_color

        if token.startswith('**') and token.endswith('**'):
            run.text = token[2:-2]
            run.bold = True
            run.font.color.rgb = RGBColor(14, 43, 100) # dark blue emphasis
        elif token.startswith('*') and token.endswith('*'):
            run.text = token[1:-1]
            run.italic = True
        elif token.startswith('`') and token.endswith('`'):
            run.text = token[1:-1]
            run.font.name = 'Consolas'
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(180, 83, 9) # amber text for inline code
            run.bold = True
        else:
            run.text = token

def build_docx(md_path, output_path):
    doc = docx.Document()

    # Define standard page setup
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Title Banner / Header
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(12)
    title_p.paragraph_format.space_after = Pt(4)
    run_t = title_p.add_run("CORREIO DIGITAL ANGOLA (CDA 2026)")
    run_t.font.name = "Calibri"
    run_t.font.size = Pt(22)
    run_t.bold = True
    run_t.font.color.rgb = RGBColor(14, 43, 100) # #0E2B64

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_p.paragraph_format.space_before = Pt(0)
    sub_p.paragraph_format.space_after = Pt(20)
    run_s = sub_p.add_run("MANUAL DE UTILIZAÇÃO OFICIAL — PLATAFORMA DE CORRESPONDÊNCIA E GOVERNAÇÃO DIGITAL")
    run_s.font.name = "Calibri"
    run_s.font.size = Pt(13)
    run_s.bold = True
    run_s.font.color.rgb = RGBColor(100, 116, 139) # slate gray

    with open(md_path, "r", encoding="utf-8") as f:
        lines = f.readlines()

    in_table = False
    table_rows = []
    in_code_block = False
    code_lines = []

    def flush_table():
        nonlocal in_table, table_rows
        if not table_rows:
            in_table = False
            return
        
        # Parse markdown tables
        # Ignore separator row (|---|---|)
        clean_rows = []
        for r in table_rows:
            if re.match(r'^\s*\|?\s*[-=:]+\s*(\|\s*[-=:]+\s*)*\|?\s*$', r):
                continue
            cells = [c.strip() for c in r.strip().strip('|').split('|')]
            clean_rows.append(cells)

        if not clean_rows:
            in_table = False
            table_rows = []
            return

        num_cols = max(len(r) for r in clean_rows)
        tbl = doc.add_table(rows=len(clean_rows), cols=num_cols)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        set_table_borders(tbl)

        for r_idx, row_data in enumerate(clean_rows):
            for c_idx in range(num_cols):
                cell = tbl.cell(r_idx, c_idx)
                set_cell_margins(cell, top=100, bottom=100, left=140, right=140)
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                cell.paragraphs[0].paragraph_format.space_before = Pt(2)
                cell.paragraphs[0].paragraph_format.space_after = Pt(2)
                
                txt = row_data[c_idx] if c_idx < len(row_data) else ""
                
                if r_idx == 0:
                    set_cell_background(cell, "0E2B64") # Navy header
                    add_formatted_text(cell.paragraphs[0], txt, default_font_size=Pt(10), default_color=RGBColor(255, 255, 255))
                    for r in cell.paragraphs[0].runs:
                        r.bold = True
                else:
                    if r_idx % 2 == 1:
                        set_cell_background(cell, "F8FAFC") # alternating light gray
                    else:
                        set_cell_background(cell, "FFFFFF")
                    add_formatted_text(cell.paragraphs[0], txt, default_font_size=Pt(9.5), default_color=RGBColor(51, 65, 85))

        doc.add_paragraph().paragraph_format.space_after = Pt(6)
        in_table = False
        table_rows = []

    def flush_code_block():
        nonlocal in_code_block, code_lines
        if not code_lines:
            in_code_block = False
            return
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(10)
        p.paragraph_format.left_indent = Inches(0.4)
        p.paragraph_format.right_indent = Inches(0.4)
        
        # Border box around code block
        pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:left w:val="single" w:sz="18" w:space="8" w:color="0E2B64"/></w:pBdr>')
        p._p.get_or_add_pPr().append(pBdr)
        shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F1F5F9"/>')
        p._p.get_or_add_pPr().append(shd)

        full_code = "".join(code_lines).strip()
        run = p.add_run(full_code)
        run.font.name = "Consolas"
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor(30, 41, 59)

        in_code_block = False
        code_lines = []

    for idx, raw_line in enumerate(lines):
        line = raw_line.rstrip('\r\n')
        
        # Check code block toggle
        if line.strip().startswith('```'):
            if in_code_block:
                flush_code_block()
            else:
                if in_table:
                    flush_table()
                in_code_block = True
            continue

        if in_code_block:
            code_lines.append(raw_line)
            continue

        # Check tables
        if line.strip().startswith('|') or (in_table and '|' in line):
            in_table = True
            table_rows.append(line)
            continue
        elif in_table:
            flush_table()

        # Skip horizontal rules or empty lines
        if line.strip() == '---' or line.strip() == '***' or line.strip() == '===':
            continue

        if not line.strip():
            continue

        # Headings
        if line.startswith('# '):
            h1 = doc.add_paragraph()
            h1.paragraph_format.space_before = Pt(18)
            h1.paragraph_format.space_after = Pt(6)
            run = h1.add_run(line[2:].strip())
            run.font.name = "Calibri"
            run.font.size = Pt(18)
            run.bold = True
            run.font.color.rgb = RGBColor(14, 43, 100)
            continue
        elif line.startswith('## '):
            h2 = doc.add_paragraph()
            h2.paragraph_format.space_before = Pt(14)
            h2.paragraph_format.space_after = Pt(4)
            run = h2.add_run(line[3:].strip())
            run.font.name = "Calibri"
            run.font.size = Pt(15)
            run.bold = True
            run.font.color.rgb = RGBColor(30, 58, 138)
            continue
        elif line.startswith('### '):
            h3 = doc.add_paragraph()
            h3.paragraph_format.space_before = Pt(10)
            h3.paragraph_format.space_after = Pt(3)
            run = h3.add_run(line[4:].strip())
            run.font.name = "Calibri"
            run.font.size = Pt(12.5)
            run.bold = True
            run.font.color.rgb = RGBColor(55, 48, 163)
            continue
        elif line.startswith('#### '):
            h4 = doc.add_paragraph()
            h4.paragraph_format.space_before = Pt(8)
            h4.paragraph_format.space_after = Pt(2)
            run = h4.add_run(line[5:].strip())
            run.font.name = "Calibri"
            run.font.size = Pt(11.5)
            run.bold = True
            run.font.color.rgb = RGBColor(71, 85, 105)
            continue

        # Bullet lists (* or -)
        if re.match(r'^\s*[\*\-]\s+', line):
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(3)
            text_part = re.sub(r'^\s*[\*\-]\s+', '', line)
            add_formatted_text(p, text_part)
            continue

        # Numbered lists (1. , 2. )
        if re.match(r'^\s*\d+\.\s+', line):
            p = doc.add_paragraph(style='List Number')
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(3)
            text_part = re.sub(r'^\s*\d+\.\s+', '', line)
            add_formatted_text(p, text_part)
            continue

        # Blockquotes (> )
        if line.startswith('> '):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.left_indent = Inches(0.4)
            pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:left w:val="single" w:sz="18" w:space="8" w:color="0E2B64"/></w:pBdr>')
            p._p.get_or_add_pPr().append(pBdr)
            shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F8FAFC"/>')
            p._p.get_or_add_pPr().append(shd)
            add_formatted_text(p, line[2:].strip(), default_font_size=Pt(10), default_color=RGBColor(71, 85, 105))
            for run in p.runs:
                run.italic = True
            continue

        # Regular Paragraph
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.15
        add_formatted_text(p, line)

    if in_table:
        flush_table()
    if in_code_block:
        flush_code_block()

    doc.save(output_path)
    print(f"Documento Word gerado com sucesso em: {output_path}")

if __name__ == "__main__":
    base_dir = "/home/user/cda-2026"
    md_file = os.path.join(base_dir, "MANUAL_DE_UTILIZACAO_CDA_2026.md")
    out_docx1 = os.path.join(base_dir, "MANUAL_DE_UTILIZACAO_CDA_2026.docx")
    out_docx2 = "/home/user/MANUAL_DE_UTILIZACAO_CDA_2026.docx"

    build_docx(md_file, out_docx1)
    build_docx(md_file, out_docx2)
